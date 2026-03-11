"""Additional test coverage for profile, resume validation, file path detection,
DOCX builder, and placeholder skip logic."""

import json
import os
import tempfile
from unittest.mock import patch

import pytest
from docx import Document

from src.models import Profile, Identity, ResumeContent
from src.profile import (
    load_profile,
    save_profile,
    create_profile,
    delete_profile,
)
from src.resume_parser import (
    _looks_like_file_path,
    _convert_windows_path,
    validate_resume_content,
)
from src.docx_builder import build_resume, _build_docx_file
from src.resume_reviewer import resolve_resume_placeholders

# ---------------------------------------------------------------------------
# Profile system: create, load, save, reset, placeholder cleanup before save
# ---------------------------------------------------------------------------


@pytest.fixture
def profile_dir(tmp_path):
    profile_path = tmp_path / "profile.json"
    with patch("src.profile.get_profile_dir", return_value=str(tmp_path)), patch(
        "src.profile.get_profile_path", return_value=str(profile_path)
    ), patch("src.profile._migrate_legacy_profile"):
        yield tmp_path, str(profile_path)


class TestProfileCreateLoadSaveReset:
    def test_create_profile_saves_to_disk(self, profile_dir, sample_resume):
        _, profile_path = profile_dir
        mock_identity = {
            "name": "Sarah Chen",
            "email": "sarah.chen@email.com",
            "phone": None,
            "location": "Seattle, WA",
            "linkedin": None,
            "github": None,
        }
        with patch("src.profile.call_llm", return_value=json.dumps(mock_identity)):
            profile = create_profile(sample_resume)

        assert os.path.isfile(profile_path)
        assert profile.identity.name == "Sarah Chen"
        assert profile.base_resume == sample_resume

    def test_load_returns_none_when_missing(self, profile_dir):
        assert load_profile() is None

    def test_save_then_load_roundtrip(self, profile_dir):
        profile = Profile(
            identity=Identity(name="Jane Doe", email="jane@test.com"),
            base_resume="Jane Doe\nSenior Engineer",
            experience_bank={"Python": "10 years"},
        )
        save_profile(profile)
        loaded = load_profile()
        assert loaded is not None
        assert loaded.identity.name == "Jane Doe"
        assert loaded.identity.email == "jane@test.com"
        assert loaded.experience_bank["Python"] == "10 years"
        assert loaded.updated_at is not None

    def test_reset_profile(self, profile_dir):
        _, profile_path = profile_dir
        profile = Profile(identity=Identity(name="Test"))
        save_profile(profile)
        assert os.path.isfile(profile_path)
        assert delete_profile() is True
        assert not os.path.isfile(profile_path)
        assert load_profile() is None

    def test_reset_nonexistent_returns_false(self, profile_dir):
        assert delete_profile() is False

    def test_placeholder_not_saved_in_base_resume(self, profile_dir):
        """Verify that saving a profile preserves text as-is; caller is
        responsible for cleaning placeholders before calling save."""
        resume_with_placeholder = "Reduced latency by [X%] through optimization"
        profile = Profile(
            identity=Identity(name="Test"),
            base_resume=resume_with_placeholder,
        )
        save_profile(profile)
        loaded = load_profile()
        # The text is stored verbatim — cleanup should happen before save
        assert loaded.base_resume == resume_with_placeholder

    def test_save_overwrites_previous(self, profile_dir):
        p1 = Profile(identity=Identity(name="First"))
        save_profile(p1)
        p2 = Profile(identity=Identity(name="Second"))
        save_profile(p2)
        loaded = load_profile()
        assert loaded.identity.name == "Second"


# ---------------------------------------------------------------------------
# Resume validation: reject non-resume content
# ---------------------------------------------------------------------------


class TestResumeValidation:
    def test_valid_resume(self, sample_resume):
        assert validate_resume_content(sample_resume) is True

    def test_rejects_random_text(self):
        assert validate_resume_content("Hello world, this is a test.") is False

    def test_rejects_empty_string(self):
        assert validate_resume_content("") is False

    def test_rejects_code_snippet(self):
        code = "def hello():\n    print('hello world')\n\nhello()"
        assert validate_resume_content(code) is False

    def test_rejects_short_gibberish(self):
        assert validate_resume_content("asdf jkl; qwerty") is False

    def test_accepts_minimal_resume(self):
        """A resume-like text with email + keyword should pass."""
        text = (
            "John Smith\n"
            "john@example.com\n"
            "555-123-4567\n"
            "Experience\n"
            "Software Engineer at Acme Corp for 5 years building distributed "
            "systems and microservices. Led a team of engineers in designing "
            "scalable infrastructure and cloud deployments across multiple regions."
        )
        assert validate_resume_content(text) is True

    def test_rejects_jd_text(self):
        """A job description without personal contact info shouldn't pass as a resume."""
        jd = "We are looking for a software engineer to join our team."
        assert validate_resume_content(jd) is False


# ---------------------------------------------------------------------------
# File path detection: Windows path to WSL conversion
# ---------------------------------------------------------------------------


class TestWindowsPathConversion:
    def test_backslash_path(self):
        result = _convert_windows_path(r"C:\Users\john\resume.docx")
        assert result == "/mnt/c/Users/john/resume.docx"

    def test_forward_slash_path(self):
        result = _convert_windows_path("D:/Documents/resume.txt")
        assert result == "/mnt/d/Documents/resume.txt"

    def test_lowercase_drive(self):
        result = _convert_windows_path(r"c:\files\resume.pdf")
        assert result == "/mnt/c/files/resume.pdf"

    def test_unix_path_unchanged(self):
        result = _convert_windows_path("/home/user/resume.txt")
        assert result == "/home/user/resume.txt"

    def test_tilde_path_unchanged(self):
        result = _convert_windows_path("~/Documents/resume.txt")
        assert result == "~/Documents/resume.txt"

    def test_mixed_separators(self):
        result = _convert_windows_path(r"E:\work/projects\resume.docx")
        assert result == "/mnt/e/work/projects/resume.docx"


class TestLooksLikeFilePathWindows:
    def test_windows_drive_letter_detected(self):
        assert _looks_like_file_path(r"C:\Users\john\resume.docx") is True

    def test_windows_forward_slash_detected(self):
        assert _looks_like_file_path("D:/Documents/resume.txt") is True

    def test_plain_text_not_detected(self):
        assert _looks_like_file_path("Jane Smith, Senior Engineer") is False


# ---------------------------------------------------------------------------
# DOCX builder: contact info from profile, categorized skills
# ---------------------------------------------------------------------------


class TestDocxBuilderContactFromProfile:
    def test_profile_identity_overrides_resume_data(self):
        """build_resume should override contact fields from Identity."""
        resume_data = ResumeContent(
            name="Original Name",
            email="original@email.com",
            phone="000-000-0000",
            summary="A summary.",
        )
        identity = Identity(
            name="Profile Name",
            email="profile@email.com",
            phone="999-999-9999",
            location="Portland, OR",
            linkedin="linkedin.com/in/profile",
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            paths = build_resume(resume_data, output_dir=tmpdir, identity=identity)
            doc = Document(paths[0])
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Profile Name" in full_text
            assert "profile@email.com" in full_text
            assert "999-999-9999" in full_text
            assert "Portland, OR" in full_text
            # Original values should be overridden
            assert "Original Name" not in full_text
            assert "original@email.com" not in full_text


class TestDocxBuilderCategorizedSkills:
    def test_categorized_skills_rendered(self):
        """Skills provided as 'Category: items' should render with bold labels."""
        resume_data = ResumeContent(
            name="Test User",
            skills=[
                "Languages: Python, Go, TypeScript",
                "Infrastructure: Kubernetes, Docker, Terraform",
                "Cloud: AWS, GCP",
            ],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Languages:" in full_text
            assert "Python, Go, TypeScript" in full_text
            assert "Infrastructure:" in full_text
            assert "Cloud:" in full_text

    def test_categorized_skills_bold_label(self):
        """The category label should be bold, items should not."""
        resume_data = ResumeContent(
            name="Test User",
            skills=["Languages: Python, Go"],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            # Find the paragraph with "Languages:"
            for para in doc.paragraphs:
                if "Languages:" in para.text:
                    # First run should be bold (the label)
                    assert para.runs[0].bold is True
                    # Second run should not be bold (the items)
                    assert para.runs[1].bold is not True
                    break
            else:
                pytest.fail("Languages: paragraph not found")

    def test_flat_skills_rendered_as_comma_list(self):
        """Skills without ':' should render as a single comma-separated paragraph."""
        resume_data = ResumeContent(
            name="Test User",
            skills=["Python", "Go", "Docker"],
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Python, Go, Docker" in full_text

    def test_skills_from_dict_normalized(self):
        """Skills dict from Claude API should be normalized to 'Category: items' format."""
        data = {
            "name": "Test User",
            "skills": {
                "Languages": ["Python", "Go"],
                "Cloud": ["AWS", "GCP"],
            },
        }
        resume_data = ResumeContent.from_dict(data)
        assert "Languages: Python, Go" in resume_data.skills
        assert "Cloud: AWS, GCP" in resume_data.skills


# ---------------------------------------------------------------------------
# Placeholder skip logic: verify no text corruption after skip
# ---------------------------------------------------------------------------


class TestPlaceholderSkipNoCorruption:
    def _skip_all(self, text):
        with patch("click.prompt", return_value="skip"), patch("click.echo"):
            return resolve_resume_placeholders(text)

    def test_skip_preserves_all_words(self):
        """Every non-placeholder word in the input should survive a skip."""
        text = "Deployed scalable infrastructure serving [X%] more traffic globally"
        result = self._skip_all(text)
        for word in (
            "Deployed",
            "scalable",
            "infrastructure",
            "serving",
            "traffic",
            "globally",
        ):
            assert word in result

    def test_skip_no_partial_word_damage(self):
        """Regression: words adjacent to removed placeholders must not lose characters."""
        text = "Migrated [number] services to Kubernetes, reducing downtime by [X%]"
        result = self._skip_all(text)
        assert "Migrated" in result
        assert "services" in result
        assert "Kubernetes" in result
        assert "reducing" in result
        assert "downtime" in result
        assert "[number]" not in result
        assert "[X%]" not in result

    def test_skip_does_not_double_space(self):
        """After removal, text should not have double spaces."""
        text = "Improved system reliability by [X%] across regions"
        result = self._skip_all(text)
        assert "  " not in result

    def test_skip_single_placeholder_sentence_intact(self):
        text = "Optimized database queries reducing latency by [X%] for all endpoints"
        result = self._skip_all(text)
        # Should still be a coherent sentence
        assert result.startswith("Optimized")
        assert "endpoints" in result
        assert "[X%]" not in result

    def test_skip_at_end_of_sentence(self):
        text = "Cut deployment time by [X%]."
        result = self._skip_all(text)
        assert "[X%]" not in result
        assert result.endswith(".")
        assert "Cut deployment time" in result


# ---------------------------------------------------------------------------
# New model fields: GapEntry.adjacent_skills, Profile new fields
# ---------------------------------------------------------------------------


class TestGapEntryAdjacentSkills:
    def test_gap_entry_has_adjacent_skills(self):
        from src.models import GapEntry

        gap = GapEntry(
            skill="Kubernetes",
            question="How many clusters have you managed?",
            adjacent_skills=["Docker", "ECS", "container orchestration"],
        )
        assert len(gap.adjacent_skills) == 3
        assert "Docker" in gap.adjacent_skills

    def test_gap_entry_default_empty_adjacent(self):
        from src.models import GapEntry

        gap = GapEntry(skill="Python", question="How many years?")
        assert gap.adjacent_skills == []

    def test_gap_analysis_from_dict_with_adjacent(self):
        from src.models import GapAnalysis

        data = {
            "gaps": [
                {
                    "skill": "Kubernetes",
                    "question": "Experience?",
                    "adjacent_skills": ["Docker", "ECS"],
                }
            ],
            "strengths": ["Python"],
        }
        result = GapAnalysis.from_dict(data)
        assert result.gaps[0].adjacent_skills == ["Docker", "ECS"]


class TestProfileWritingPreferences:
    def test_profile_to_dict_includes_new_fields(self):
        profile = Profile(
            identity=Identity(name="Test"),
            base_resume="resume",
            original_resume="original",
            writing_preferences={"tone": "formal"},
            applications_since_review=3,
        )
        d = profile.to_dict()
        assert d["original_resume"] == "original"
        assert d["writing_preferences"] == {"tone": "formal"}
        assert d["applications_since_review"] == 3

    def test_profile_from_dict_with_new_fields(self):
        data = {
            "identity": {"name": "Test"},
            "base_resume": "resume",
            "original_resume": "original",
            "writing_preferences": {"bullet_length": "shorter"},
            "applications_since_review": 7,
        }
        profile = Profile.from_dict(data)
        assert profile.original_resume == "original"
        assert profile.writing_preferences["bullet_length"] == "shorter"
        assert profile.applications_since_review == 7

    def test_profile_from_dict_defaults(self):
        """Old profiles without new fields should get defaults."""
        data = {
            "identity": {"name": "Test"},
            "base_resume": "resume",
        }
        profile = Profile.from_dict(data)
        assert profile.original_resume == ""
        assert profile.writing_preferences == {}
        assert profile.applications_since_review == 0
