"""Tests for docx_builder.py."""

import os
import tempfile

import pytest
from docx import Document

from src.docx_builder import build_resume, _build_docx_file, _build_markdown
from src.models import ResumeContent


@pytest.fixture
def resume_data(mock_resume_generation):
    """Convert fixture dict to ResumeContent model."""
    return ResumeContent.from_dict(mock_resume_generation)


class TestBuildDocxFile:
    def test_creates_docx(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)
            assert os.path.exists(filepath)

    def test_docx_has_name(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            texts = [p.text for p in doc.paragraphs]
            assert any("Sarah Chen" in t for t in texts)

    def test_docx_has_contact_info(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "sarah.chen@email.com" in full_text
            assert "(415) 555-0198" in full_text
            assert "Seattle, WA" in full_text

    def test_docx_has_sections(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for heading in (
                "SUMMARY",
                "EXPERIENCE",
                "SKILLS",
                "EDUCATION",
                "CERTIFICATIONS",
            ):
                assert heading in full_text

    def test_docx_has_experience_bullets(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "2M+ events/hour" in full_text
            assert "CloudScale Technologies" in full_text

    def test_docx_has_education(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "University of Washington" in full_text
            assert "M.S. Computer Science" in full_text

    def test_docx_has_certifications(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(resume_data, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "AWS Certified Solutions Architect" in full_text

    def test_docx_contact_from_profile_override(self, mock_resume_generation):
        """Contact info override works via model attributes."""
        data = dict(mock_resume_generation)
        data["email"] = "override@newmail.com"
        data["phone"] = "(999) 000-0000"
        rd = ResumeContent.from_dict(data)

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.docx")
            _build_docx_file(rd, filepath)

            doc = Document(filepath)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "override@newmail.com" in full_text
            assert "(999) 000-0000" in full_text


class TestBuildMarkdown:
    def test_creates_md_file(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.md")
            _build_markdown(resume_data, filepath)
            assert os.path.exists(filepath)

    def test_md_has_name_header(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.md")
            _build_markdown(resume_data, filepath)

            with open(filepath) as f:
                content = f.read()
            assert "# Sarah Chen" in content

    def test_md_has_sections(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.md")
            _build_markdown(resume_data, filepath)

            with open(filepath) as f:
                content = f.read()
            for heading in (
                "## Professional Summary",
                "## Experience",
                "## Skills",
                "## Education",
                "## Certifications",
            ):
                assert heading in content

    def test_md_has_bullet_points(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = os.path.join(tmpdir, "test.md")
            _build_markdown(resume_data, filepath)

            with open(filepath) as f:
                content = f.read()
            assert "- Architected real-time event processing" in content


class TestBuildResume:
    def test_default_format_docx(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            paths = build_resume(resume_data, output_dir=tmpdir)
            assert len(paths) == 1
            assert paths[0].endswith(".docx")
            assert os.path.exists(paths[0])

    def test_md_format(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            paths = build_resume(resume_data, output_dir=tmpdir, formats=["md"])
            assert len(paths) == 1
            assert paths[0].endswith(".md")

    def test_output_path_as_directory(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            outdir = os.path.join(tmpdir, "custom_output")
            os.makedirs(outdir)
            paths = build_resume(resume_data, output_path=outdir)
            assert all(outdir in p for p in paths)

    def test_output_path_as_filename(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            outpath = os.path.join(tmpdir, "my_resume.docx")
            paths = build_resume(resume_data, output_path=outpath)
            assert len(paths) == 1
            assert "my_resume" in os.path.basename(paths[0])

    def test_all_formats_expands(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            # Only test docx and md — pdf requires libreoffice
            paths = build_resume(resume_data, output_dir=tmpdir, formats=["docx", "md"])
            extensions = {os.path.splitext(p)[1] for p in paths}
            assert ".docx" in extensions
            assert ".md" in extensions

    def test_creates_output_dir_if_missing(self, resume_data):
        with tempfile.TemporaryDirectory() as tmpdir:
            outdir = os.path.join(tmpdir, "new", "nested", "dir")
            paths = build_resume(resume_data, output_dir=outdir)
            assert os.path.isdir(outdir)
            assert len(paths) == 1
