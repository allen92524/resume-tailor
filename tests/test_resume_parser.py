"""Tests for resume_parser.py."""

import os
import tempfile

import pytest

from unittest.mock import patch

from src.resume_parser import (
    read_resume_from_file,
    _looks_like_file_path,
    _convert_docker_path,
)

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


class TestLooksLikeFilePath:
    def test_absolute_path(self):
        assert _looks_like_file_path("/home/user/resume.txt") is True

    def test_tilde_path(self):
        assert _looks_like_file_path("~/Documents/resume.txt") is True

    def test_pasted_text(self):
        assert _looks_like_file_path("John Smith") is False

    def test_multiline_text(self):
        assert _looks_like_file_path("John Smith\nemail@example.com") is False

    def test_empty_string(self):
        assert _looks_like_file_path("") is False

    def test_whitespace_before_path(self):
        # _looks_like_file_path strips leading whitespace, so this is detected as a path
        assert _looks_like_file_path("  /home/user/resume.txt") is True

    def test_relative_path_not_detected(self):
        assert _looks_like_file_path("resume.txt") is False


class TestReadResumeFromFile:
    def test_read_txt(self):
        path = os.path.join(FIXTURES_DIR, "sample_resume.txt")
        text = read_resume_from_file(path)
        assert "Sarah Chen" in text
        assert "CloudScale Technologies" in text

    def test_read_md(self):
        with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False) as f:
            f.write("# Jane Doe\n\n## Experience\n- Software Engineer at Acme")
            f.flush()
            try:
                text = read_resume_from_file(f.name)
                assert "Jane Doe" in text
                assert "Software Engineer" in text
            finally:
                os.unlink(f.name)

    def test_read_docx(self):
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Test Resume Name")
            doc.add_paragraph("Software Engineer at TestCorp")
            doc.save(f.name)
            try:
                text = read_resume_from_file(f.name)
                assert "Test Resume Name" in text
                assert "TestCorp" in text
            finally:
                os.unlink(f.name)

    def test_read_pdf(self):
        pytest.importorskip("reportlab")
        from io import BytesIO
        from reportlab.pdfgen import canvas as rl_canvas

        buf = BytesIO()
        c = rl_canvas.Canvas(buf)
        c.drawString(72, 720, "PDF Resume Content")
        c.save()
        buf.seek(0)

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(buf.read())
            f.flush()
            try:
                text = read_resume_from_file(f.name)
                assert "PDF Resume Content" in text
            finally:
                os.unlink(f.name)

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            read_resume_from_file("/nonexistent/path/resume.txt")

    def test_unsupported_format(self):
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            f.write(b"fake")
            f.flush()
            try:
                with pytest.raises(ValueError, match="Unsupported file format"):
                    read_resume_from_file(f.name)
            finally:
                os.unlink(f.name)

    def test_tilde_expansion(self):
        # Should not crash even if path doesn't exist after expansion
        with pytest.raises(FileNotFoundError):
            read_resume_from_file("~/nonexistent_resume_12345.txt")


class TestConvertDockerPath:
    """Test Docker host-path to container-mount-path conversion."""

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_macos_downloads(self, _mock):
        assert _convert_docker_path("/Users/jane/Downloads/resume.pdf") == "/mnt/downloads/resume.pdf"

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_linux_desktop(self, _mock):
        assert _convert_docker_path("/home/jane/Desktop/resume.docx") == "/mnt/desktop/resume.docx"

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_linux_documents(self, _mock):
        assert _convert_docker_path("/home/user/Documents/cv.txt") == "/mnt/documents/cv.txt"

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_windows_path_after_wsl_conversion(self, _mock):
        # Windows paths are first converted to /mnt/c/... by _convert_windows_path
        assert _convert_docker_path("/mnt/c/Users/Jane/Downloads/resume.pdf") == "/mnt/downloads/resume.pdf"

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_tilde_downloads(self, _mock):
        assert _convert_docker_path("~/Downloads/resume.pdf") == "/mnt/downloads/resume.pdf"

    @patch("src.resume_parser._is_docker", return_value=True)
    def test_unmatched_path_unchanged(self, _mock):
        assert _convert_docker_path("/tmp/resume.pdf") == "/tmp/resume.pdf"

    @patch("src.resume_parser._is_docker", return_value=False)
    def test_not_in_docker_unchanged(self, _mock):
        assert _convert_docker_path("/Users/jane/Downloads/resume.pdf") == "/Users/jane/Downloads/resume.pdf"


