"""Build formatted resume output from structured data."""

import logging
import os
import platform
import re
import subprocess
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import Identity, JDAnalysis, ResumeContent

logger = logging.getLogger(__name__)

# Common abbreviations for role titles
_ROLE_ABBREVIATIONS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\bSoftware Engineer\b", re.I), "SWE"),
    (re.compile(r"\bSoftware Engineering\b", re.I), "SWE"),
    (re.compile(r"\bSenior\b", re.I), "Sr"),
    (re.compile(r"\bJunior\b", re.I), "Jr"),
    (re.compile(r"\bInfrastructure\b", re.I), "Infra"),
    (re.compile(r"\bEngineering\b", re.I), "Eng"),
    (re.compile(r"\bEngineer\b", re.I), "Eng"),
    (re.compile(r"\bManager\b", re.I), "Mgr"),
    (re.compile(r"\bManagement\b", re.I), "Mgmt"),
    (re.compile(r"\bDevelopment\b", re.I), "Dev"),
    (re.compile(r"\bDeveloper\b", re.I), "Dev"),
    (re.compile(r"\bPrincipal\b", re.I), "Principal"),
    (re.compile(r"\bArchitect\b", re.I), "Arch"),
    (re.compile(r"\bAdministrator\b", re.I), "Admin"),
    (re.compile(r"\bDirector\b", re.I), "Dir"),
    (re.compile(r"\bVice President\b", re.I), "VP"),
    (re.compile(r"\bAssistant\b", re.I), "Asst"),
    (re.compile(r"\bAssociate\b", re.I), "Assoc"),
    (re.compile(r"\bCoordinator\b", re.I), "Coord"),
    (re.compile(r"\bSpecialist\b", re.I), "Spec"),
    (re.compile(r"\bConsultant\b", re.I), "Consult"),
    (re.compile(r"\bRepresentative\b", re.I), "Rep"),
    (re.compile(r"\bSupervisor\b", re.I), "Supv"),
    (re.compile(r"\bTechnician\b", re.I), "Tech"),
    (re.compile(r"\bProfessor\b", re.I), "Prof"),
    (re.compile(r"\bLieutenant\b", re.I), "Lt"),
    (re.compile(r"\bSergeant\b", re.I), "Sgt"),
]


def _make_output_basename(
    identity: Identity | None,
    jd_analysis: JDAnalysis | None,
) -> str:
    """Build a descriptive filename like Name_Company_Role from available data.

    Falls back to timestamped name if identity/JD info is missing.
    """
    parts: list[str] = []

    # Name from identity
    if identity and identity.name:
        name = identity.name.strip()
        # Take first and last name
        name_parts = name.split()
        if len(name_parts) >= 2:
            parts.append(f"{name_parts[0]}_{name_parts[-1]}")
        else:
            parts.append(name_parts[0])

    # Company from JD analysis
    if jd_analysis and jd_analysis.company:
        parts.append(jd_analysis.company.strip())

    # Role from JD analysis (abbreviated)
    if jd_analysis and jd_analysis.job_title:
        role = jd_analysis.job_title.strip()
        for pattern, abbrev in _ROLE_ABBREVIATIONS:
            role = pattern.sub(abbrev, role)
        parts.append(role)

    if not parts:
        return f"resume_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"

    # Join, replace spaces with underscores, remove special chars
    basename = "_".join(parts)
    basename = basename.replace(" ", "_")
    basename = re.sub(r"[^\w\-]", "", basename)
    # Collapse multiple underscores
    basename = re.sub(r"_+", "_", basename).strip("_")
    return basename


def _set_run_font(
    run,
    name: str = "Calibri",
    size: float = 10,
    bold: bool = False,
    color: str | None = None,
):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_bullet_paragraph(doc, text: str):
    """Add a paragraph formatted as a bullet point with a round bullet character."""
    para = doc.add_paragraph()
    run = para.add_run(f"\u2022  {text}")
    _set_run_font(run, size=10.5)
    pf = para.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(3)
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.15)
    pf.line_spacing = 1.15
    return para


def build_resume(
    resume_data: ResumeContent,
    output_dir: str = "output",
    output_path: str | None = None,
    formats: list[str] | None = None,
    identity: Identity | None = None,
    jd_analysis: JDAnalysis | None = None,
) -> list[str]:
    """Build resume output in the requested formats.

    Args:
        resume_data: Structured ResumeContent data.
        output_dir: Default output directory (used when output_path is None).
        output_path: Explicit output path. If a directory, files go there.
            If a file path, used as the base name.
        formats: List of formats to generate. Options: "docx", "pdf", "md", "all".
            Defaults to ["docx"].
        identity: Profile identity for contact info and filename.
        jd_analysis: JD analysis for filename (company/role).

    Returns:
        List of generated file paths.
    """
    if formats is None:
        formats = ["docx"]
    if "all" in formats:
        formats = ["docx", "pdf", "md"]

    # Override contact info with profile identity (source of truth)
    if identity:
        for field_name in ("name", "email", "phone", "location", "linkedin"):
            profile_value = getattr(identity, field_name, None)
            if profile_value:
                setattr(resume_data, field_name, profile_value)

    # Determine output directory and base filename
    base_name = _make_output_basename(identity, jd_analysis)

    if output_path:
        output_path = os.path.expanduser(output_path)
        if os.path.isdir(output_path) or output_path.endswith("/"):
            dest_dir = output_path
        else:
            dest_dir = os.path.dirname(output_path) or "."
            # Use the provided filename as base (strip extension)
            base_name = os.path.splitext(os.path.basename(output_path))[0]
    else:
        dest_dir = output_dir

    os.makedirs(dest_dir, exist_ok=True)

    generated: list[str] = []

    # Always build the DOCX first (needed for PDF conversion too)
    needs_docx = "docx" in formats or "pdf" in formats
    docx_path = os.path.join(dest_dir, f"{base_name}.docx")

    if needs_docx:
        _build_docx_file(resume_data, docx_path)
        if "docx" in formats:
            generated.append(os.path.abspath(docx_path))

    if "pdf" in formats:
        pdf_path = os.path.join(dest_dir, f"{base_name}.pdf")
        _convert_docx_to_pdf(docx_path, pdf_path)
        generated.append(os.path.abspath(pdf_path))
        # Clean up intermediate DOCX if not requested
        if "docx" not in formats and os.path.exists(docx_path):
            os.remove(docx_path)

    if "md" in formats:
        md_path = os.path.join(dest_dir, f"{base_name}.md")
        _build_markdown(resume_data, md_path)
        generated.append(os.path.abspath(md_path))

    logger.info("Built %d output file(s): %s", len(generated), generated)
    return generated



def open_file(filepath: str) -> None:
    """Open a file with the system default application."""
    system = platform.system()
    # Check for WSL
    is_wsl = False
    if system == "Linux":
        try:
            with open("/proc/version", "r") as f:
                is_wsl = "microsoft" in f.read().lower()
        except OSError:
            pass

    try:
        if is_wsl:
            subprocess.Popen(
                ["wslview", filepath],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        elif system == "Linux":
            subprocess.Popen(
                ["xdg-open", filepath],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        elif system == "Darwin":
            subprocess.Popen(
                ["open", filepath], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
        elif system == "Windows":
            os.startfile(filepath)
        else:
            logger.warning("Don't know how to open files on %s", system)
    except FileNotFoundError:
        logger.warning("Could not find a file opener. Please open the file manually.")


# --- Private builders ---


def _build_docx_file(resume_data: ResumeContent, filepath: str) -> None:
    """Build and save a DOCX file."""
    logger.debug("Building DOCX: %s", filepath)
    doc = Document()

    # Set default font to Calibri for clean PDF rendering
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Set page margins: 0.5" top/bottom, 0.6" left/right
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- Header: Name (18pt, bold, centered, dark) ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_para.paragraph_format.space_before = Pt(0)
    name_run = name_para.add_run(resume_data.name)
    _set_run_font(name_run, size=14, bold=True, color="1A1A1A")

    # --- Header: Contact info (10pt, centered, gray, pipe-separated) ---
    contact_parts: list[str] = []
    for field in ("email", "phone", "location", "linkedin"):
        value = getattr(resume_data, field)
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(4)
        contact_para.paragraph_format.space_before = Pt(0)
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        _set_run_font(contact_run, size=10.5, color="555555")

    # --- Summary (10pt, normal paragraph) ---
    if resume_data.summary:
        _add_section_heading(doc, "Summary")
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(4)
        summary_para.paragraph_format.space_before = Pt(0)
        summary_para.paragraph_format.line_spacing = 1.15
        summary_run = summary_para.add_run(resume_data.summary)
        _set_run_font(summary_run, size=10.5)

    # --- Experience ---
    if resume_data.experience:
        _add_section_heading(doc, "Experience")
        for i, job in enumerate(resume_data.experience):
            # Job title | Company | Dates — all on one line
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(4) if i > 0 else Pt(0)
            title_para.paragraph_format.space_after = Pt(2)

            # Title in bold
            title_run = title_para.add_run(job.title)
            _set_run_font(title_run, size=10.5, bold=True, color="1A1A1A")

            # Company (normal weight)
            company_run = title_para.add_run(f" | {job.company}")
            _set_run_font(company_run, size=10.5, color="1A1A1A")

            # Dates (normal weight)
            if job.dates:
                dates_run = title_para.add_run(f" | {job.dates}")
                _set_run_font(dates_run, size=10.5, color="1A1A1A")

            # Bullets
            for bullet in job.bullets:
                _add_bullet_paragraph(doc, bullet)

    # --- Skills (10pt, grouped with bold category labels) ---
    if resume_data.skills:
        _add_section_heading(doc, "Skills")
        # Check if skills are formatted as "Category: items" or plain list
        categorized = [s for s in resume_data.skills if ":" in s]
        if categorized:
            # Render each categorized skill line with bold label
            for skill_line in resume_data.skills:
                skill_para = doc.add_paragraph()
                skill_para.paragraph_format.space_after = Pt(2)
                skill_para.paragraph_format.space_before = Pt(0)
                skill_para.paragraph_format.line_spacing = 1.15
                if ":" in skill_line:
                    label, items = skill_line.split(":", 1)
                    label_run = skill_para.add_run(f"{label.strip()}:")
                    _set_run_font(label_run, size=10.5, bold=True)
                    items_run = skill_para.add_run(f" {items.strip()}")
                    _set_run_font(items_run, size=10.5)
                else:
                    run = skill_para.add_run(skill_line)
                    _set_run_font(run, size=10.5)
        else:
            # Plain comma-separated list
            skills_para = doc.add_paragraph()
            skills_para.paragraph_format.space_after = Pt(2)
            skills_para.paragraph_format.line_spacing = 1.15
            skills_run = skills_para.add_run(", ".join(resume_data.skills))
            _set_run_font(skills_run, size=10.5)

    # --- Education ---
    if resume_data.education:
        _add_section_heading(doc, "Education")
        for edu in resume_data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(2)
            edu_para.paragraph_format.space_before = Pt(0)
            edu_para.paragraph_format.line_spacing = 1.15

            # Degree in bold
            edu_run = edu_para.add_run(edu.degree)
            _set_run_font(edu_run, size=10.5, bold=True)

            if edu.institution:
                inst_run = edu_para.add_run(f" | {edu.institution}")
                _set_run_font(inst_run, size=10.5)
            if edu.year:
                year_run = edu_para.add_run(f" | {edu.year}")
                _set_run_font(year_run, size=10.5)

    # --- Certifications ---
    if resume_data.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in resume_data.certifications:
            _add_bullet_paragraph(doc, cert)

    # --- Licenses ---
    if resume_data.licenses:
        _add_section_heading(doc, "Licenses")
        for lic in resume_data.licenses:
            _add_bullet_paragraph(doc, lic)

    # --- Publications ---
    if resume_data.publications:
        _add_section_heading(doc, "Publications")
        for pub in resume_data.publications:
            _add_bullet_paragraph(doc, pub)

    # --- Awards ---
    if resume_data.awards:
        _add_section_heading(doc, "Awards")
        for award in resume_data.awards:
            _add_bullet_paragraph(doc, award)

    # --- Volunteer ---
    if resume_data.volunteer:
        _add_section_heading(doc, "Volunteer Experience")
        for vol in resume_data.volunteer:
            _add_bullet_paragraph(doc, vol)

    doc.save(filepath)
    logger.debug("DOCX saved: %s", filepath)


def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """Convert a DOCX file to PDF using LibreOffice."""
    output_dir = os.path.dirname(pdf_path) or "."
    logger.info("Converting DOCX to PDF: %s -> %s", docx_path, pdf_path)
    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
        # LibreOffice names the output based on input filename
        lo_output = os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        if lo_output != pdf_path and os.path.exists(lo_output):
            os.rename(lo_output, pdf_path)
    except FileNotFoundError:
        raise RuntimeError(
            "PDF conversion requires LibreOffice. "
            "Install it with: sudo apt install libreoffice-writer"
        )


def _build_markdown(resume_data: ResumeContent, filepath: str) -> None:
    """Build a Markdown version of the resume."""
    logger.debug("Building Markdown: %s", filepath)
    lines: list[str] = []

    lines.append(f"# {resume_data.name}")
    lines.append("")

    contact_parts: list[str] = []
    for field_name in ("email", "phone", "location", "linkedin"):
        value = getattr(resume_data, field_name)
        if value:
            contact_parts.append(value)
    if contact_parts:
        lines.append(" | ".join(contact_parts))
        lines.append("")

    if resume_data.summary:
        lines.append("## Professional Summary")
        lines.append("")
        lines.append(resume_data.summary)
        lines.append("")

    if resume_data.experience:
        lines.append("## Experience")
        lines.append("")
        for job in resume_data.experience:
            header = f"**{job.title} — {job.company}**"
            if job.dates:
                header += f"  |  {job.dates}"
            lines.append(header)
            lines.append("")
            for bullet in job.bullets:
                lines.append(f"- {bullet}")
            lines.append("")

    if resume_data.skills:
        lines.append("## Skills")
        lines.append("")
        for skill_line in resume_data.skills:
            lines.append(skill_line)
        lines.append("")

    if resume_data.education:
        lines.append("## Education")
        lines.append("")
        for edu in resume_data.education:
            edu_text = edu.degree
            if edu.institution:
                edu_text += f" — {edu.institution}"
            if edu.year:
                edu_text += f" ({edu.year})"
            lines.append(f"- {edu_text}")
        lines.append("")

    if resume_data.certifications:
        lines.append("## Certifications")
        lines.append("")
        for cert in resume_data.certifications:
            lines.append(f"- {cert}")
        lines.append("")

    if resume_data.licenses:
        lines.append("## Licenses")
        lines.append("")
        for lic in resume_data.licenses:
            lines.append(f"- {lic}")
        lines.append("")

    if resume_data.publications:
        lines.append("## Publications")
        lines.append("")
        for pub in resume_data.publications:
            lines.append(f"- {pub}")
        lines.append("")

    if resume_data.awards:
        lines.append("## Awards")
        lines.append("")
        for award in resume_data.awards:
            lines.append(f"- {award}")
        lines.append("")

    if resume_data.volunteer:
        lines.append("## Volunteer Experience")
        lines.append("")
        for vol in resume_data.volunteer:
            lines.append(f"- {vol}")
        lines.append("")

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def _add_section_heading(doc: Document, title: str):
    """Add a section heading: 11pt, bold, uppercase, dark blue, with thin line underneath."""
    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run(title.upper())
    _set_run_font(heading_run, size=11, bold=True, color="2B379C")

    pf = heading_para.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(8)

    # Thin bottom border in dark blue (ATS-safe approach)
    p_pr = heading_para._element.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B379C")
    p_borders.append(bottom)
    p_pr.append(p_borders)
